# -*- coding: utf8 -*-

import urllib2,json,time,re
from docx import Document
from docx.shared import Inches
import sys,socket,pickle,os


#http://bbs.tianya.cn/post-no05-148332-1.shtml 测试地址



all_urls_dict = {}
def get_all_urls(public_next_id,tech_next_id,city_next_id):

    url_interface = "http://www.tianya.cn/api/tw?method=userinfo.ice.getUserTotalArticleList&" \
                    "params.userId=2213624&" \
                    "params.pageSize=20&" \
                    "params.bMore=true&" \
                    "params.publicNextId=%s&" \
                    "params.techNextId=%s&" \
                    "params.cityNextId=%s" %(public_next_id,tech_next_id,city_next_id)

    headers = {'Referer':'http://www.tianya.cn/2213624/bbs?t=post','user-agent':'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.103 Safari/537.36'}
    url = urllib2.Request(url_interface,headers=headers)
    respone = urllib2.urlopen(url)
    data_dict = json.load(respone)

    if data_dict['success'] == u'1':
        data = data_dict['data']
        rows = data_dict['data']['rows']
        public_next_id = data['public_next_id']
        tech_next_id = data['tech_next_id']
        city_next_id = data['city_next_id']
        for row in rows:
            title = row['title']
            art_id = row['art_id']
            item = row['item']
            url = 'http://bbs.tianya.cn/post-%s-%s-1.shtml' %(item,art_id)
            all_urls_dict[title.encode('UTF-8')] = url.encode('UTF-8')
        if len(rows) == 20:
            get_all_urls(public_next_id,tech_next_id,city_next_id)
    else:
        print '获取接口数据失败.'






def get_all_page(index_url):
    "获取此贴的所有页面保存到本地"
    global headers
    headers = {'Referer':index_url,'User-Agent':'Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.103 Safari/537.36'}


    try:
        req = urllib2.Request(index_url, headers=headers)
        html = urllib2.urlopen(req,timeout=120).read()
        if html is None:sys.exit('获取不到首页内容')
        #获取基本信息
        global host,page_len,title,d_value
        d_value = set([])
        host = getHost(html)
        page_len = getPageLen(html)
        title = getTitle(html)

        print '基础信息获取成功...'

        html_list = os.listdir('./html')
        if len(html_list) != page_len:
            local_html_set = set([int(html[:-6]) for html in html_list])
            all_html_set = set(range(1,page_len+1))
            d_value =  all_html_set - local_html_set

            for i in d_value:
                try:
                    url = re.sub(r'-(\d{1,4})\.','-'+str(i)+'.',index_url)
                    req = urllib2.Request(url,headers=headers)
                    html = urllib2.urlopen(req).read()
                    file_name = './html/%d.shtml' %i
                    with open(file_name,'wb') as f:
                        f.write(html)
                    time.sleep(5)
                    print 'download %s success.' %url
                except urllib2.URLError:
                    print 'get %s raise a error' %url
                except IOError:
                    print 'write the %s raise a error' %url

            local_html_set = set([int(html[:-6]) for html in html_list])
            all_html_set = set(range(1,page_len+1))
            d_value =  all_html_set - local_html_set


    except urllib2.URLError,e:
        print 'URLError'
        print e
    except socket.error:
        print 'socket error...'



def getMainContent():
    "获取主贴内容，默认为本地1开头的第一个shtml"
    #if len(d_value) == 0 : return

    with open('./html/1.shtml','rb') as f:
        html = f.read()
    res = re.findall(r'<div class="bbs-content clearfix">(.*?)</div>',html,re.S)

    if not len(res):
        sys.exit('楼主主贴没有获取到内容...')
    elif len(res) > 1:
        sys.eixt('楼主主贴获取的内容超过一个...')
    global _main_body
    with open('./txt/mainContent.txt','wb') as f:
        #f.write(formatHtml(res[0]))
        res = formatHtml(res[0])
        f.write(res)



def getComment():
    "获取楼主的评论"
    with open('./txt/comment.txt','wb+') as f2:
        f2.write('--------评论部分--------\n')
        for i in range(1,page_len+1):
            file_name = './html/%d.shtml' %i
            with open(file_name,'rb') as f:
                html = f.read()
            comment_regx = r'<div class="atl-item" _host="%s".*?<div class="bbs-content">(.*?)</div>' %host
            res = re.findall(comment_regx,html,re.S)

            if len(res):
                for comment in res:
                    comment = formatHtml(comment)
                    f2.write(comment)
                    f2.write('\n')




def saveImg(img_url):
    "保存图片在当前项目下，默认图片名为demo"
    form = img_url[-3:]

    try:
        req = urllib2.Request(img_url,headers=headers)
        rsp = urllib2.urlopen(req).read()
        img_name = 'demo.'+form
        with open('./img/'+img_name,'wb') as f:
            f.write(rsp)
        return img_name
    except urllib2.URLError,e:
        print img_url
        print e
        return 0

def save2Word():
    file_name = './word/'+title+'.docx'


    with open('./txt/mainContent.txt','a') as f1:
        with open('./txt/comment.txt','rb') as f2:
            f1.write(f2.read())

    with open('./txt/mainContent.txt','rb') as article_complete:

        document = Document()
        for line in article_complete:
            if len(line) == 7 or len(line) == 13 or line == '\n':continue #手动换行符暂时解决不了

            img_url = re.findall(r'(http:.*?\.jpg|http:.*?\.png|http:.*?\.gif)',line,re.S)
            line = re.sub(r'(http:.*?\.jpg|http:.*?\.png|http:.*?\.gif)','',line)
            for img in img_url:
                imgName = saveImg(img)
                if imgName != 0:
                    try:
                        document.add_picture('./img/'+imgName, width=Inches(3.25))
                    except Exception:
                         document.add_paragraph(u'图片加载失败....')
                         document.add_paragraph(img)
                else:
                    document.add_paragraph(u'图片加载失败....')
                    document.add_paragraph(img)
            document.add_paragraph(line.decode('UTF-8'))

        document.save(file_name.decode('UTF-8'))


def formatHtml(line):

    line = line.replace('<br>','\n') #去掉<br>标签
    line = re.sub(r'<img src="http://.*?original="','',line)
    line = line.replace('<img src="http://img3','http://img3')
    line = line.replace('" />','')
    line = line.replace('">','')
    line = line.replace('&lt;','<') # 替换小于号
    line = line.replace('&gt;','>') #替换大于号
    line = line.replace('&quot;','"') #替换双引号

    line = line.replace('\t','') #去掉\t
    line = line.strip()
    return line


def getHost(html):
    "获取楼主host"
    res = re.findall(r'<div class="atl-item host-item" _host="(.*?)">',html,re.S)
    if not len(res):
        sys.exit('获取不到楼主的host')
    elif len(res) > 1:
        sys.exit('楼主的host不是唯一')
    return res[0]

def getPageLen(html):
    #获取贴子的总页数
    res = re.findall(r'return goPage\(this,.*,(.*?)\);">',html,re.S)
    if not len(res):
        sys.exit('没有获取到总页数')
    elif len(res) > 1:
        sys.exit('获取到%s个总页面'%len(res))
    return int(res[0])

def getTitle(html):
    title = re.findall(r'<h1 class="atl-title">(.*?)</h1>',html,re.S)
    if len(title) != 1:sys.exit('获取标题有误。。。')
    title = re.sub(r'<[^>]+>','',title[0]).strip()
    return title

def save_url_to_txt(txt ='./all_urls.txt'):
    "保存所有链接到TXT里"
    get_all_urls('343675','2147483647','2147483647')
    with open(txt,'wb') as f:

        for key in all_urls_dict.keys():
            f.write(key)
            f.write(u'    ')
            f.write(all_urls_dict[key])
            f.write(u'\n')


def run():
    "运行脚本"
    print '请输入要获取的天涯贴子的第一页...'
    while True:
        input_url = raw_input()
        input_url = input_url.strip().lower()
        if input_url == 'q':sys.exit('退出成功.')
        result = re.findall(r'^http://.*html$',input_url)
        if len(result) != 1:
            print '地址格式不正确，请重新输入...或按 q 退出.'
        else:
            break

    result =raw_input('是否完全重新开始:y/n,按q退出程序\n')

    while True:
        if result not in 'ynq' or result == '':
            print '输入有误，请重新输入：y/n,按q退出程序\n'
            result =raw_input('是否完全重新开始:y/n,按q退出程序\n')
        else:
            break
    if result == 'y':
        dir_list = os.listdir('./')
        if 'html' not in dir_list:
            os.mkdir('./html')
        else:
            html_list = os.listdir('./html')
            for html in html_list:
                os.remove('./html/'+html)

    elif result == 'q':
        sys.exit('退出成功！')

    get_all_page(input_url)
    print '下载所有页面成功！'
    getMainContent()
    print '获取主贴内容成功！'
    getComment()
    print '获取楼主评论成功！'
    save2Word()
    print '完成！'


if __name__ == '__main__':
    run()
